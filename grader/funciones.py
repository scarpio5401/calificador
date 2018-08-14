#!/usr/bin/env python
# -*- coding: utf-8 -*-

import random
import zipfile
import os
import shutil
import json
import tempfile
from lxml import etree
from docx import Document
from utils import *
from django.utils.encoding import smart_str,smart_unicode

NAMESPACE = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'} #define NAMESPACE
#settings = open("settings.dat",'r').read()
#data = json.loads(settings)
data = {
	"tabla_contenido":{
		"generar":5,"estilos":10,"multinivel":10
	},
	"formato":{
		"wordart":3,"interlineado":5,"espaciado":5,"fuente":4,"tamano":3,"color":3
	},
	"nota_pie":{
		"puntaje":5
	},
	"letra_capital":[ 
		{"0-0.2":0},{"0.21-0.50":4},{"0.51-1":7}
	],
	"pie_pagina":[
		{"0-0.39":0},{"0.40-0.80":5},{"0.81-1":8}
	],
	"salto_pagina":[
		{"0-0.19":0},{"0.20-0.39":1},{"0.40-0.89":3},{"0.90-1":5}
	],
	"vineta":[
		{"0-0.19":0},{"0.20-0.39":3},{"0.40-0.89":7},{"0.90-1":10}
	],
	"columna":[
		{"0-0.39":0},{"0.40-0.89":7},{"0.9-1":10}
	],
	"bordes":{
		"color":3,"contorno":4,"grosor":3
	}
}

#Devuelve el object element de una archivo xml, arg1 archivo, arg2 tipo [0:document.xml,1:footnotes.xml,2:footer1.xml]
def get_xml(documento,tipo):
   hash = str(random.randint(1000,9999))
   tmp_path = tempfile.gettempdir()
   doc_path = tmp_path+'/'+hash+'.docx'
   #copia = open(doc_path,'w+')
   #shutil.copyfileobj(documento,copia)
   copia = Document(documento)
   copia.save(doc_path)
   #shutil.copy(documento,hash+'.docx')
   copy_path = tmp_path+'/'+hash
   os.rename(doc_path,copy_path+'.zip')
   fzip=zipfile.ZipFile(copy_path+'.zip','r')
   fzip.extractall(copy_path)
   fzip.close()
   if tipo==0: doc = copy_path+"/word/document.xml"
   if tipo==1: doc = copy_path+"/word/footnotes.xml"
   if tipo==2: doc = copy_path+"/word/footer1.xml"
   if tipo==3: doc = copy_path+"/word/footer2.xml"
   if tipo==4: doc = copy_path+"/word/footer3.xml"
   if not os.path.isfile(doc): return None
   xml = open(doc,'r')
   tree = etree.fromstring(xml.read())
   xml.close()
   os.remove(copy_path+'.zip')
   shutil.rmtree(copy_path)
   return tree

   
def get_xml_footers(documento):
   arr=[]
   hash = str(random.randint(1000,9999))
   tmp_path = tempfile.gettempdir()
   doc_path = tmp_path+'/'+hash+'.docx'
   copia = Document(documento)
   copia.save(doc_path)
   copy_path = tmp_path+'/'+hash
   os.rename(doc_path,copy_path+'.zip')
   
   fzip=zipfile.ZipFile(copy_path+'.zip','r')
   fzip.extractall(copy_path)
   fzip.close()
   docs = os.listdir(copy_path+"/word")
   for d in docs:
      if d.startswith('footer'):
	     xml = open(copy_path+"/word/"+d,'r')
	     tree = etree.fromstring(xml.read())
	     arr.append(tree)
	     xml.close()
   os.remove(copy_path+'.zip')
   shutil.rmtree(copy_path)
   return arr

def validar_tc(modelo,respuesta):
   #Verificar si genero tabla de contenido
   links_respuesta = get_links_tc(respuesta.element)
   links_modelo = get_links_tc(modelo.element)
   porcentaje=comparar_arreglos(links_modelo,links_respuesta)
   nota_generar = round(data['tabla_contenido']['generar']*porcentaje)
   #Verificar titulos con estilo
   titulos_respuesta = get_paras_estilo(respuesta)
   titulos_modelo = get_titulos(modelo)
   porcentaje=comparar_arreglos(titulos_modelo,titulos_respuesta)
   nota_titulos = round(data['tabla_contenido']['estilos']*porcentaje)
   #Verficar titulos con lista multinivel
   titulos_respuesta = get_titulos(respuesta)
   porcentaje=comparar_arreglos(titulos_modelo,titulos_respuesta)
   #Si no tiene estilo titulo, verificar numeración
   if  porcentaje==0:
      porcentaje=tiene_numeracion(respuesta,titulos_modelo)
   nota_multinivel = round(data['tabla_contenido']['multinivel']*porcentaje)
   #print str(nota_generar)+"-"+str(nota_titulos)+"-"+str(nota_multinivel)
   return [nota_generar,nota_titulos,nota_multinivel]  

#Devuelve arreglo con los links de una tabla de contenido   
def get_links_tc(documento):
   links = []
   body = documento.find('w:body', NAMESPACE)
   tc = body.find('w:sdt', NAMESPACE)
   if tc is not None:
      content_tc = tc.find('w:sdtContent', NAMESPACE)
      if content_tc is not None:
         hlinks = content_tc.findall(".//w:hyperlink", NAMESPACE)
         for l in hlinks:
	        texto = l.findall('.//w:t',NAMESPACE)
	        links.append(get_texto(texto))
   #print links
   return links

#Devuelve arreglo con los titulos de un documento 
def get_titulos(documento):
   arr = []
   for para in documento.paragraphs:
      texto = smart_unicode(para.text)
      if texto.strip()!='':
	     style = para.style
	     #print smart_str(para.text[:20])
	     #print tiene_estilo(para.style)
	     if es_estilo_titulo(style.name): 
		    arr.append(texto)
	     else:
	        base = style.base_style
	        if base is not None:
	           if es_estilo_titulo(base.name): 
			      arr.append(texto)
   return arr

#Devuelve arreglo con los titulos de un documento 
def get_paras_estilo(documento):
   arr = []
   for para in documento.paragraphs:
      texto = smart_unicode(para.text)
      if texto.strip()!='':
	     if tiene_estilo(para.style):
	        arr.append(texto)
   return arr   
   
def es_estilo_titulo(stylename):
   if "Título" in str(stylename): return True
   if "Heading" in str(stylename): return True
   return False

def tiene_estilo(style):
   if es_estilo_titulo(style.name):return True
   base = style.base_style
   if base:
      if base.name=="Normal":
         return False
      else:
         return True
   else:
      return False

def tiene_numeracion(documento,titulos):
   n=len(titulos)
   i=float(0)
   porcentaje = 0
   paras =  documento.paragraphs
   for t in titulos:
      for p in paras:
         if t==p.text:
		    if p.style and p.style.name:
		       if p.style.name=='List Paragraph':i=i+0.5 	
		       if p.style.name=='header':i=i+1 			   
   porcentaje = float(i)/float(n)
   porcentaje = 1 if porcentaje>1 else porcentaje
   return porcentaje
   
def validar_notapie(doc_modelo,doc_respuesta):
   notas_modelo = get_notapie(doc_modelo)
   notas_respuesta = get_notapie(doc_respuesta)
   porcentaje = comparar_arreglos(notas_modelo,notas_respuesta)
   nota_notapie = round(data['nota_pie']['puntaje']*porcentaje)
   return nota_notapie
      
def get_notapie(documento):
   arr=[]
   xml = get_xml(documento,1)
   if xml is None: return arr
   notas = xml.findall("w:footnote", NAMESPACE)
   for n in notas:
      nota = n.findall(".//w:t", NAMESPACE)
      if len(nota)>0:
	     str=''
	     for e in nota:
	        str=str+e.text
	     arr.append(str)
   return arr

def validar_letracap(modelo,respuesta):
   letracap_modelo = get_letracap(modelo)
   letracap_respuesta = get_letracap(respuesta)
   porcentaje = comparar_arreglos(letracap_modelo,letracap_respuesta)
   rangos = data['letra_capital']
   return get_puntaje_rango(rangos,porcentaje)

def get_letracap(documento):
   arr = []
   body = documento.find('w:body', NAMESPACE)
   capitales = body.findall(".//w:framePr", NAMESPACE)
   if capitales is not None:
      for c in capitales:
         capital = c.getparent().getparent().find(".//w:t", NAMESPACE)
         if capital is not None:
            arr.append(capital.text)
   return arr

def validar_saltos(modelo, respuesta):
   saltos_modelo = get_saltos(modelo)
   saltos_respuesta = get_saltos(respuesta)
   #print "mod:"+str(saltos_modelo)+" res:"+str(saltos_respuesta)
   saltos_respuesta = saltos_modelo if saltos_respuesta>saltos_modelo else saltos_respuesta
   porcentaje = float(saltos_respuesta)/float(saltos_modelo)
   rangos = data['salto_pagina']
   return get_puntaje_rango(rangos,porcentaje)

def get_saltos(body):
   cont=0
   saltos = body.findall(".//w:br[@w:type='page']", NAMESPACE)
   if saltos is not None:
      for s in saltos:
         cont=cont+1
   return cont

def validar_vinetas(modelo, respuesta):
   vinetas_modelo = get_vinetas(modelo)
   vinetas_respuesta = get_vinetas(respuesta)
   #print vinetas_modelo
   #print vinetas_respuesta
   porcentaje = comparar_arreglos(vinetas_modelo,vinetas_respuesta)
   rangos = data['vineta']
   return get_puntaje_rango(rangos,porcentaje)   
   
def get_vinetas(body):
   arr = []
   vinetas = body.findall(".//w:numPr", NAMESPACE)
   if vinetas is not None:
      for v in vinetas:
	     pstyle = v.getparent().findall(".//w:pStyle[@w:val='Prrafodelista']",NAMESPACE)
	     if len(pstyle)>0:
		    p = v.getparent().getparent()
		    texto= p.findall(".//w:t", NAMESPACE)
		    if len(texto)>0:
		       str=''
		       for e in texto:
		          str=str+e.text
		       arr.append(str)
   return arr   

def validar_columnas(modelo, respuesta):
   columnas_modelo = get_columnas(modelo)
   columnas_respuesta = get_columnas(respuesta)
   porcentaje = comparar_arreglos(columnas_modelo,columnas_respuesta)
   rangos = data['columna']
   return get_puntaje_rango(rangos,porcentaje)  
   
def get_columnas(body):
   arr = []
   num,sep = '', ''
   columnas = body.findall(".//w:sectPr", NAMESPACE)
   if columnas is not None:
      for c in columnas:
         columna = c.find('.//w:cols[@w:num="2"]', NAMESPACE)
         if columna is not None:
			values = columna.values()
			if len(values)>1:
				num = values[0]
				sep = values[1]
   return [num, sep]

def validar_piepagina(doc_modelo,doc_respuesta):
   piepagina_modelo = get_piepagina(doc_modelo)
   piepagina_respuesta = get_piepagina(doc_respuesta)
   #print "modelo"+str(piepagina_modelo)
   #print "respuesta"+str(piepagina_respuesta)
   porcentaje = comparar_arreglos(piepagina_modelo,piepagina_respuesta)
   #print porcentaje
   rangos = data['pie_pagina']
   nota_piepagina = get_puntaje_rango(rangos,porcentaje)  
   return nota_piepagina
   
def get_piepagina(documento):
   arr=[]
   xmls = get_xml_footers(documento)
   for xml in xmls:
	  texto=xml.findall(".//w:t",NAMESPACE)
	  num_pag=xml.findall(".//w:instrText",NAMESPACE)
	  if len(texto)>0:
	     for e in texto:
	        arr.append(e.text)
	  if len(num_pag)>0:
	     for e in num_pag:
	        arr.append(e.text)
   return arr
   
def validar_formato(doc_modelo,doc_respuesta):
   arr_resp=[0,0,0,0,0,0]
   doc_mod = doc_modelo #cargar_documet(doc_modelo)
   doc_res = doc_respuesta #cargar_documet(doc_respuesta)
   paras_mod = cargar_paras(doc_mod)
   paras_res = cargar_paras(doc_res)
   num_paras = len(paras_mod)
   for para in paras_res:
	  indice = buscar_para(para,paras_mod)
	  #print smart_str(para[0][:20])
	  if indice is not None: 
	     p_mod_obj = doc_mod.paragraphs[indice]
	     p_res_obj = doc_res.paragraphs[para[1]]
	     arr_nota = val_format(p_mod_obj,p_res_obj)
	     arr_resp = sumar_arrs(arr_resp,arr_nota)
   porcentajes = [x/num_paras for x in arr_resp]
   fuente = calc_nota(data['formato']['fuente'],porcentajes[0])
   color = calc_nota(data['formato']['color'],porcentajes[1])
   tamano = calc_nota(data['formato']['tamano'],porcentajes[2])
   espaciado = calc_nota(data['formato']['espaciado'],((porcentajes[3]+porcentajes[4])/2))
   interl = calc_nota(data['formato']['interlineado'],porcentajes[5])
   return [interl,espaciado,fuente,tamano,color]

def calc_nota(puntaje,porcentaje):
   if porcentaje>=0.85: # 85% de aciertos en el formato
      porcentaje = 1
   nota = round(puntaje*porcentaje)
   return nota
   
def cargar_documet(ruta):
   f = open(ruta,'rb')
   document = Document(f)
   f.close()
   return document

def cargar_paras(document):
   arr = []
   index = 0
   paras = document.paragraphs
   for p in paras:
      texto = p.text.strip()
      if texto!='' and len(texto)>1:
         arr.append([p.text,index])
      index+=1
   return arr
   
def buscar_para(para, arr_paras):
   texto = para[0]
   for p in arr_paras:
      if comparar_cad(texto,p[0]):
	     return p[1]
   return None
		 
def val_format(p_mod_obj,p_res_obj):
   #print "R:f:"+str(get_fuente(p_res_obj))+",c:"+str(get_color(p_res_obj))+",t:"+str(get_tamano(p_res_obj))+",i:"+str(get_interlin(p_res_obj))+",p:"+str(get_espacpos(p_res_obj))+",a:"+str(get_espacant(p_res_obj))
   #print "M:f:"+str(get_fuente(p_mod_obj))+",c:"+str(get_color(p_mod_obj))+",t:"+str(get_tamano(p_mod_obj))+",i:"+str(get_interlin(p_mod_obj))+",p:"+str(get_espacpos(p_mod_obj))+",a:"+str(get_espacant(p_mod_obj))
   fuente = 1 if get_fuente(p_mod_obj)==get_fuente(p_res_obj) else 0
   color  = 1 if get_color(p_mod_obj)==get_color(p_res_obj) else 0
   tamano = 1 if get_tamano(p_mod_obj)==get_tamano(p_res_obj) else 0
   esppos = 1 if get_espacpos(p_mod_obj)==get_espacpos(p_res_obj) else 0
   espant = 1 if get_espacant(p_mod_obj)==get_espacant(p_res_obj) else 0
   interl = 1 if get_interlin(p_mod_obj)==get_interlin(p_res_obj) else 0
   return [fuente,color,tamano,esppos,espant,interl]

def get_fuente(para_obj):
   style = para_obj.style
   if style.font and style.font.name: return style.font.name
   base = style.base_style
   if base and base.font and base.font.name: return base.font.name
   runs = para_obj.runs
   if runs and runs[0] and runs[0].font: return para_obj.runs[0].font.name
   return None

def get_color(para_obj):
   style = para_obj.style
   if style.font and style.font.color and style.font.color.rgb: 
      return style.font.color.rgb
   base = style.base_style
   if base and base.font and base.font.color and base.font.color.rgb: return base.font.color.rgb
   runs = para_obj.runs
   if runs and runs[0].font and runs[0].font.color: return para_obj.runs[0].font.color.rgb
   return None

def get_tamano(para_obj):
   style = para_obj.style
   if style.font and style.font.size and style.font.size.pt: 
      return style.font.size.pt
   base = style.base_style
   if base:
      if base.font and base.font.size: return base.font.size.pt
   runs = para_obj.runs
   if runs and runs[0].font and runs[0].font.size: return para_obj.runs[0].font.size.pt
   return None

def get_espacpos(para_obj):
   style = para_obj.style
   para_format = para_obj.paragraph_format
   if para_format and para_format.space_after and para_format.space_after.pt: 
      return para_format.space_after.pt
   if style.paragraph_format and style.paragraph_format.space_after:
      return style.paragraph_format.space_after.pt
   return None

def get_espacant(para_obj):
   style = para_obj.style
   para_format = para_obj.paragraph_format
   if para_format and para_format.space_before and para_format.space_before.pt: 
      return para_format.space_before.pt
   if style.paragraph_format and style.paragraph_format.space_before:
      return style.paragraph_format.space_before.pt
   return None

def get_interlin(para_obj):
   style = para_obj.style
   para_format = para_obj.paragraph_format
   if para_format and para_format.line_spacing: return para_format.line_spacing
   if style.paragraph_format and style.paragraph_format.line_spacing:
      return style.paragraph_format.line_spacing
   return None

def validar_bordes(doc_modelo,doc_respuesta):
   nota_bordes = [0,0,0] #color,contorno,grosor
   estilo_modelo = cargar_estilo_bordes(doc_modelo)
   num_paras = 3 #por revisar
   respuesta = doc_respuesta#cargar_documet(doc_respuesta)
   para_bordes = get_paras_bordes(respuesta)
   for b in para_bordes:
      nota_para = get_nota_para(b,estilo_modelo)
      nota_bordes = sumar_arrs(nota_bordes,nota_para)
   porcentajes = [x/num_paras for x in nota_bordes]
   nota_col = round(data['bordes']['color']*porcentajes[0])
   nota_con = round(data['bordes']['contorno']*porcentajes[1])
   nota_gro = round(data['bordes']['grosor']*porcentajes[2])
   return [nota_col,nota_con,nota_gro]

def cargar_estilo_bordes(document):
   #document = cargar_documet(ruta)
   para = document.paragraphs[0]
   border = para._element.findall(".//w:pBdr", NAMESPACE)
   estilo = get_estilo_borde(border[0],'top')
   return estilo
   
def get_estilo_borde(element,posicion):
   border = element.find('w:'+posicion,NAMESPACE)
   if border is not None:
      values = border.values()
   else:
      return [None,None,None]
   return [values[3],values[0],values[1]] #color,contorno,grosor

def get_paras_bordes(document):
   arr = []
   paras = document.paragraphs
   for p in paras:
      texto = p.text.strip()
      if texto!='' and len(texto)>1:
	     border = p._element.findall(".//w:pBdr", NAMESPACE)
	     if len(border)>0 and border[0] is not None:
		    arr.append(border[0])
   return arr   
   
def get_nota_para(border,estilo):
   nota = [0,0,0]
   estilo_top = get_estilo_borde(border,'top')
   estilo_bottom = get_estilo_borde(border,'bottom')
   estilo_left = get_estilo_borde(border,'left')
   estilo_right = get_estilo_borde(border,'right')
   nota_top = comparar_estilos(estilo_top,estilo)
   nota_bot = comparar_estilos(estilo_bottom,estilo)
   nota_lef = comparar_estilos(estilo_left,estilo)
   nota_rig = comparar_estilos(estilo_right,estilo)
   nota[0] = (nota_top[0]+nota_bot[0]+nota_lef[0]+nota_rig[0])/float(4) #color
   nota[1] = (nota_top[1]+nota_bot[1]+nota_lef[1]+nota_rig[1])/float(4) #contorno
   nota[2] = (nota_top[2]+nota_bot[2]+nota_lef[2]+nota_rig[2])/float(4) #grosor
   return nota
   
def comparar_estilos(arr1,arr2):
   result = [0,0,0]
   n = len(arr2)
   for i in range(n):
	  #si color es auto igual a negro(000000)
	  if arr2[i]=='auto': 
		 arr2[i]='000000' 
	  if arr1[i]=='auto': 
	     arr1[i]='000000' 
	  if arr1[i]==arr2[i]:
	     result[i]=1
	  else: 
	     result[i]=0
   return result